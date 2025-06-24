import os
import PyPDF2
import docx
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from io import BytesIO
from utils.text_utils import anonymize_text_content

class ResumeProcessor:
    """Service for processing and extracting information from resumes"""
    
    def __init__(self):
        """Initialize the resume processor"""
        pass
    
    def extract_text(self, resume_path):
        """
        Extract text from resume file (PDF, DOCX, or TXT)
        
        Args:
            resume_path (str): Path to the resume file
            
        Returns:
            str: Extracted text from the resume, or an error message string.
        """
        if not os.path.exists(resume_path):
            return f"Error: File not found at {resume_path}"
            
        file_extension = os.path.splitext(resume_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(resume_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_from_docx(resume_path)
            elif file_extension == '.txt':
                return self._extract_from_txt(resume_path)
            else:
                return f"Error: Unsupported file format: {file_extension}"
        except Exception as e:
            return f"Error extracting text from {os.path.basename(resume_path)}: {str(e)}"

    def _extract_from_pdf(self, pdf_path):
        """Extract text from PDF file"""
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                if not pdf_reader.pages:
                    return "[PDF is empty or unreadable]\n"
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        # This can happen with image-based PDFs
                        text += f"[PDF page {page_num+1} has no extractable text]\n"
            if not text.strip():
                 return "[PDF contains no extractable text]\n"
        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"
        return text
    
    def _extract_from_docx(self, docx_path):
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = docx.Document(docx_path)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            if not text.strip():
                return "[DOCX contains no extractable text]\n"
        except Exception as e:
            if isinstance(e, docx.opc.exceptions.PackageNotFoundError):
                 return f"Error: Could not open file. It might be a legacy .doc file or corrupted: {str(e)}"
            return f"Error extracting DOCX text: {str(e)}"
        return text
    
    def _extract_from_txt(self, txt_path):
        """Extract text from TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            if not text.strip():
                return "[TXT file is empty]\n"
            return text
        except Exception as e:
            return f"Error extracting TXT text: {str(e)}"
    
    def _add_paragraph_with_style(self, doc, text, style_name=None, font_size=11, bold=False, space_after=Pt(6), space_before=Pt(0), level=0, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, font_name='Calibri'):
        p = doc.add_paragraph(text, style=style_name)
        p.paragraph_format.space_after = space_after
        p.paragraph_format.space_before = space_before
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.alignment = alignment
        
        if style_name == 'List Bullet':
            p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
            p.paragraph_format.first_line_indent = Inches(-0.25)
        else:
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)

        for run in p.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
        return p

    def generate_one_pager_docx(self, resume_path: str) -> BytesIO:
        """
        Extracts text from a resume, anonymizes it, and generates a 
        well-formatted, organized DOCX file in memory by parsing Markdown-like output.
        """
        raw_text = self.extract_text(resume_path)
        
        if raw_text.startswith("Error:") or "[PDF" in raw_text or "[DOCX" in raw_text or "[TXT" in raw_text:
            raise ValueError(f"Could not process resume content: {raw_text}")

        anonymized_text = anonymize_text_content(raw_text) # This calls the LLM

        if anonymized_text.startswith("Error during anonymization"):
             raise ValueError(anonymized_text)

        doc = DocxDocument()
        # Set document margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        lines = anonymized_text.split('\n')
        
        first_line_processed = False

        for line in lines:
            stripped_line = line.strip()
            
            if not stripped_line: # Handle genuinely empty lines from LLM for spacing
                doc.add_paragraph().paragraph_format.space_after = Pt(3) # Minimal space
                continue
            
            # Skip lines that are just the REDACTED placeholder
            if stripped_line == "[REDACTED]":
                continue

            # Handle Markdown headings
            if stripped_line.startswith('# ') and not first_line_processed: # Level 1 heading for First Name
                text_content = stripped_line[2:].strip()
                if text_content: # Ensure there's actual text after the hash
                    self._add_paragraph_with_style(doc, text_content, font_size=18, bold=True, space_after=Pt(12), space_before=Pt(0), alignment=WD_ALIGN_PARAGRAPH.CENTER)
                    first_line_processed = True
            elif stripped_line.startswith('## '): # Level 2 heading for main sections
                text_content = stripped_line[3:].strip()
                if text_content:
                    self._add_paragraph_with_style(doc, text_content, font_size=14, bold=True, space_after=Pt(6), space_before=Pt(10), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            elif stripped_line.startswith('### '): # Level 3 heading for sub-sections/titles
                text_content = stripped_line[4:].strip()
                if text_content:
                    self._add_paragraph_with_style(doc, text_content, font_size=12, bold=True, space_after=Pt(4), space_before=Pt(6), alignment=WD_ALIGN_PARAGRAPH.LEFT)
            # Handle Markdown bullet points
            elif stripped_line.startswith('- '):
                text_content = stripped_line[2:].strip()
                if text_content:
                    self._add_paragraph_with_style(doc, text_content, style_name='List Bullet', font_size=11, space_after=Pt(2), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, level=0)
            elif stripped_line.startswith('  - '): # Nested bullet
                text_content = stripped_line[4:].strip()
                if text_content:
                    self._add_paragraph_with_style(doc, text_content, style_name='List Bullet', font_size=11, space_after=Pt(2), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, level=1)
            elif stripped_line.startswith('    - '): # Even deeper nested bullet
                text_content = stripped_line[6:].strip()
                if text_content:
                    self._add_paragraph_with_style(doc, text_content, style_name='List Bullet', font_size=11, space_after=Pt(2), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, level=2)
            else:
                # Default to a regular paragraph with justified alignment
                # Only add if it's not the first name line again (if LLM somehow repeats it without #)
                if not (not first_line_processed and len(stripped_line.split()) <= 2 and "[REDACTED]" not in stripped_line):
                     self._add_paragraph_with_style(doc, stripped_line, font_size=11, space_after=Pt(3), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream
    
    # extract_sections is part of the original screener, no changes needed here for 1-pager
    def extract_sections(self, resume_text):
        sections = {
            'contact_info': '', 'education': '', 'experience': '',
            'skills': '', 'projects': '', 'certifications': '', 'other': ''
        }
        lines = resume_text.split('\n')
        current_section = 'other'
        for line in lines:
            stripped_line = line.strip()
            lower_line = stripped_line.lower()
            if not stripped_line: continue

            section_map = {
                'education': ['education', 'academic', 'degree'],
                'experience': ['experience', 'employment', 'work history', 'career'],
                'skills': ['skill', 'technical', 'technologies', 'languages', 'expertise'],
                'projects': ['project', 'portfolio'],
                'certifications': ['certification', 'certificate', 'license', 'training'],
                'contact_info': ['contact', 'email', 'phone', 'address', 'linkedin', 'github']
            }
            
            found_section = False
            for section_name, keywords in section_map.items():
                if any(keyword in lower_line for keyword in keywords) and len(stripped_line) < 50: # Heuristic for section headers
                    current_section = section_name
                    sections[current_section] += stripped_line + '\n'
                    found_section = True
                    break
            if not found_section:
                sections[current_section] += stripped_line + '\n'
        return sections
